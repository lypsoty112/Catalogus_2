import re
from typing import Dict, Generator, List

import pandas as pd
import docx


def parse_documents(documents: list) -> Generator[int, None, pd.DataFrame]:
    objects = []
    for i, doc in enumerate(documents):
        text = parse_document(doc)
        objects.extend(text)
        yield i
    yield pd.DataFrame.from_records(objects)


def parse_document(document) -> List[dict]:
    doc = docx.Document(document)

    title = doc.paragraphs[0].text
    full_text = ""
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            full_text += txt.lower() + " "

    # Count the amount of time LANDO, KVANTO, and DIAMETRO regex patterns are found
    object_count = max([len(re.findall(r"diametro:\s", full_text.lower())),
                        len(re.findall(r"lando:\s", full_text.lower())),
                        len(re.findall(r"kvanto:\s", full_text.lower()))])

    return [{**parse_object_lines(doc.paragraphs), "title": title}]


def parse_object_lines(lines: List) -> dict:
    return_dict: Dict[str, str | int | None] = {
        "text": "",
        "LANDO": "",
        "KVANTO": "",
        "METALO": "",
        "PEZO": "",
        "PRODUKTORO": "",
        "MEZUROJ": "",
        "DIAMETRO": "",
        "MEDALISTO": "",
        "JARO": "",
        "DIKO": ""
    }
    r_pattern = (r"(?=\s*(kvanto|metalo|pezo|produktoro|mezuroj|diametro|medalisto|jaro|diko|lando|averso|artisto"
                 r"|materialo|substanco)\s*|$)")

    regex_patterns = {
        "LANDO": rf"lando:\s*(.+?){r_pattern}",
        "KVANTO": rf"kvanto:\s*(.+?){r_pattern}",
        "METALO": rf"metalo:\s*(.+?){r_pattern}",
        "PEZO": rf"pezo:\s*(.+?){r_pattern}",
        "PRODUKTORO": rf"produktoro:\s*(.+?){r_pattern}",
        "MEZUROJ": rf"mezuroj:\s*(.+?){r_pattern}",
        "DIAMETRO": rf"diametro:\s*(.+?){r_pattern}",
        "MEDALISTO": rf"medalisto:\s*(.+?){r_pattern}",
        "JARO": rf"jaro:\s*(.+?){r_pattern}",
        "DIKO": rf"diko:\s*(.+?){r_pattern}",
        "MATERIALO": rf"materialo:\s*(.+?){r_pattern}",
        "SUBSTANCO": rf"substanco:\s*(.+?){r_pattern}",

    }
    for para in lines:
        text = para.text.strip()
        if text:
            return_dict["text"] += text + " "

    text = return_dict["text"].lower()
    for key, pattern in regex_patterns.items():
        match = re.search(pattern, text)
        if match:
            return_dict[key] = match.group(1)

    # Remove any 'nekonata' values
    unknown_values = ["nekonata", "serĉata", "?", ""]
    for key, value in return_dict.items():
        if value.lower() in unknown_values:
            return_dict[key] = None

    return return_dict
